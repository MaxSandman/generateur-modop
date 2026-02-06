import streamlit as st
import google.generativeai as genai
import cv2
import os
import tempfile
import time
import re
import zipfile
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

# --- CONFIGURATION PAGE ---
st.set_page_config(page_title="Nomadia SmartDoc", page_icon="⚡", layout="wide")

# --- DESIGN SYSTEM NOMADIA ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap');
    .stApp { background-color: #F8FAFC; font-family: 'Inter', sans-serif; }
    
    .nomadia-header { 
        text-align: center; font-size: 40px; font-weight: 800; color: #0B192E; 
        margin-top: 20px; text-transform: uppercase; letter-spacing: -1px;
    }
    .highlight { color: #A3E671; }
    .subtitle { color: #64748B; text-align: center; font-size: 16px; margin-bottom: 30px; }

    /* WORKFLOW STEPPER */
    .stepper-container {
        display: flex; justify-content: space-between; margin-bottom: 40px; 
        padding: 0 50px; position: relative;
    }
    .step {
        background: white; border: 2px solid #E2E8F0; color: #64748B;
        padding: 10px 20px; border-radius: 50px; font-weight: bold; font-size: 14px;
        z-index: 2; width: 30%; text-align: center;
    }
    .step.active {
        border-color: #A3E671; background-color: #F0FDF4; color: #0B192E; box-shadow: 0 4px 6px rgba(163, 230, 113, 0.1);
    }
    .step-line {
        position: absolute; top: 50%; left: 10%; right: 10%; height: 2px; background: #E2E8F0; z-index: 1; transform: translateY(-50%);
    }

    /* ZONES */
    .zone-card {
        background-color: white; border-radius: 15px; padding: 30px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.03); border: 1px solid #E2E8F0; margin-bottom: 20px;
    }
    .zone-title {
        font-size: 18px; font-weight: 700; color: #0B192E; margin-bottom: 15px; border-bottom: 2px solid #A3E671; display: inline-block; padding-bottom: 5px;
    }

    .summary-box {
        background-color: #F0F9FF; border-left: 5px solid #00D2B4; padding: 20px; border-radius: 8px; color: #0B192E; margin-bottom: 25px;
    }

    .stButton>button { 
        background-color: #0B192E !important; color: #A3E671 !important; 
        border: none !important; border-radius: 8px !important; 
        padding: 12px 20px !important; font-weight: bold !important; width: 100%;
    }
    </style>
    """, unsafe_allow_html=True)

# --- FONCTION EXTRACTION IMAGE ---
def extract_frame(video_path, timestamp_str):
    try:
        ts = timestamp_str.replace('[','').replace(']','').strip()
        parts = list(map(int, ts.split(':')))
        seconds = parts[0] * 60 + parts[1] if len(parts) == 2 else parts[0]
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened(): return None
        fps = cap.get(cv2.CAP_PROP_FPS)
        cap.set(cv2.CAP_PROP_POS_FRAMES, int(seconds * fps))
        ret, frame = cap.read()
        if ret:
            tfile = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
            cv2.imwrite(tfile.name, frame)
            cap.release()
            return tfile.name
        cap.release()
    except: return None
    return None

# --- CONFIGURATION API ---
api_key = st.sidebar.text_input("🔑 Clé API", type="password", value=st.secrets.get("GEMINI_API_KEY", ""))
if api_key: 
    genai.configure(api_key=api_key)
else: 
    st.info("Veuillez saisir votre clé API Google Gemini pour continuer.")
    st.stop()

# --- HEADER & STEPPER ---
st.markdown('<div class="nomadia-header">NOMADIA <span class="highlight">SMARTDOC</span></div>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Générateur de documentation technique automatisé</p>', unsafe_allow_html=True)

if 'steps' in st.session_state:
    current_step = 3
elif 'processing' in st.session_state:
    current_step = 2
else:
    current_step = 1

st.markdown(f"""
    <div class="stepper-container">
        <div class="step-line"></div>
        <div class="step {'active' if current_step >= 1 else ''}">1. IMPORT VIDÉO</div>
        <div class="step {'active' if current_step >= 2 else ''}">2. ANALYSE IA</div>
        <div class="step {'active' if current_step >= 3 else ''}">3. RÉSULTAT & EXPORT</div>
    </div>
""", unsafe_allow_html=True)

# --- ZONE 1 : IMPORT ---
if 'steps' not in st.session_state:
    st.markdown('<div class="zone-card"><div class="zone-title">📡 Source Média</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload", type=['mp4', 'mov'], label_visibility="collapsed")
    
    if uploaded_file:
        with st.expander("👁️ Voir la vidéo source"):
            st.video(uploaded_file)
        
        col_c, col_btn, col_d = st.columns([1, 2, 1])
        with col_btn:
            if st.button("LANCER L'ANALYSE INTELLIGENTE"):
                st.session_state.processing = True
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# --- ZONE 2 : TRAITEMENT ---
if 'processing' in st.session_state and 'steps' not in st.session_state:
    st.markdown('<div class="zone-card"><div class="zone-title">🧠 Traitement en cours</div>', unsafe_allow_html=True)
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.markdown("**Préparation de la vidéo...**")
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tfile:
            tfile.write(uploaded_file.read())
            video_path = tfile.name
        
        myfile = genai.upload_file(path=video_path)
        progress_bar.progress(20)
        
        while myfile.state.name == "PROCESSING":
            time.sleep(2)
            myfile = genai.get_file(myfile.name)
        progress_bar.progress(50)
        
        # --- DÉTECTION DYNAMIQUE DU MODÈLE (FIX 404) ---
        status_text.markdown("**Sélection du meilleur moteur IA...**")
        
        # On liste tous les modèles disponibles pour votre clé
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        
        # On cherche d'abord la version standard, sinon n'importe quel flash
        target_model = None
        for m in available_models:
            if "gemini-1.5-flash" in m:
                target_model = m
                break
        
        if not target_model:
            for m in available_models:
                if "flash" in m:
                    target_model = m
                    break
        
        if not target_model:
            target_model = available_models[0] # Fallback ultime

        model = genai.GenerativeModel(target_model)
        
        prompt = """Tu es un expert technique chez Nomadia. Analyse cette vidéo démonstrative.
        1. Commence impérativement par un bloc RESUME: [Décris ici la procédure globale en 3 phrases max].
        2. Détaille ensuite chaque étape avec ce format strict :
        TITRE: [Nom de l'action]
        DESC: [Explication précise]
        TIME: [MM:SS]
        Séparateur: ---"""
        
        response = model.generate_content([prompt, myfile])
        progress_bar.progress(80)
        
        # Parsing Résumé & Étapes
        summary_match = re.search(r"RESUME: (.*?)(?=---)", response.text, re.DOTALL)
        st.session_state.summary = summary_match.group(1).strip() if summary_match else "Résumé non détecté."
        
        steps_list = []
        for block in response.text.split('---'):
            t = re.search(r"TITRE: (.*)", block)
            d = re.search(r"DESC: (.*)", block)
            ts = re.search(r"TIME: (.*)", block)
            if t and d and ts:
                steps_list.append({
                    "title": t.group(1).strip(), 
                    "desc": d.group(1).strip(), 
                    "time": ts.group(1).strip(), 
                    "img": extract_frame(video_path, ts.group(1))
                })
        
        st.session_state.steps = steps_list
        progress_bar.progress(100)
        del st.session_state.processing
        st.rerun()

    except Exception as e:
        st.error(f"Erreur d'analyse : {str(e)}")
        if 'processing' in st.session_state:
            del st.session_state.processing

# --- ZONE 3 : RÉSULTATS & EXPORT ---
if 'steps' in st.session_state:
    st.markdown('<div class="zone-card">', unsafe_allow_html=True)
    st.markdown('<div class="zone-title">📋 Synthèse de la procédure</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="summary-box"><strong>📌 Résumé IA :</strong><br>{st.session_state.summary}</div>', unsafe_allow_html=True)
    
    # Chemins temporaires
    doc_path = "Nomadia_Guide.docx"
    zip_path = "Nomadia_Confluence.zip"
    pdf_path = "Nomadia_Guide.pdf"

    # Export Word
    doc = Document()
    doc.add_heading('Nomadia SmartDoc', 0)
    doc.add_heading('Résumé de la démonstration', level=1)
    doc.add_paragraph(st.session_state.summary)
    for s in st.session_state.steps:
        doc.add_heading(f"{s['title']} ({s['time']})", level=2)
        doc.add_paragraph(s['desc'])
        if s['img'] and os.path.exists(s['img']):
            try: doc.add_picture(s['img'], width=Inches(5))
            except: pass
    doc.save(doc_path)

    # Export Confluence (ZIP)
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        txt_content = f"h1. Guide d'utilisation Nomadia\n\n{{note}}{st.session_state.summary}{{note}}\n\n"
        for i, s in enumerate(st.session_state.steps):
            img_name = f"etape_{i+1}.jpg"
            txt_content += f"h2. {s['title']} ({s['time']})\n!{img_name}|width=500!\n{s['desc']}\n\n"
            if s['img'] and os.path.exists(s['img']):
                zipf.write(s['img'], arcname=img_name)
        zipf.writestr("guide_confluence.txt", txt_content)

    # Export PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "NOMADIA SMARTDOC", ln=True, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 8, f"SYNTHÈSE :\n{st.session_state.summary}")
    pdf.ln(10)
    for i, s in enumerate(st.session_state.steps):
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, f"{i+1}. {s['title']} ({s['time']})", ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 8, s['desc'])
        if s['img'] and os.path.exists(s['img']):
            try: pdf.image(s['img'], x=15, w=170)
            except: pass
        pdf.ln(5)
    pdf.output(pdf_path)

    st.markdown('<div class="zone-title">💾 Exporter le résultat</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1: 
        with open(doc_path, "rb") as f: st.download_button("📘 Télécharger Word", f, doc_path)
    with c2: 
        with open(zip_path, "rb") as f: st.download_button("🗂️ Pack Confluence (.zip)", f, zip_path)
    with c3: 
        with open(pdf_path, "rb") as f: st.download_button("📕 Télécharger PDF", f, pdf_path)
    
    st.markdown('</div>', unsafe_allow_html=True)

    with st.expander("🔎 Voir le détail des étapes et images"):
        for step in st.session_state.steps:
            img_col, txt_col = st.columns([0.3, 0.7])
            with img_col: 
                if step['img'] and os.path.exists(step['img']): st.image(step['img'])
            with txt_col: st.markdown(f"**{step['title']}** ({step['time']})\n\n{step['desc']}")
            st.divider()
    
    if st.button("🔄 Analyser une autre vidéo"):
        for k in ['steps','processing','summary']: 
            if k in st.session_state: del st.session_state[k]
        st.rerun()
